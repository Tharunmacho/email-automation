"""Extract machine-readable text from any supported resume file.

Dispatch is by detected file type, never by filename. For PDFs the embedded text
layer is used wherever it is trustworthy and every other page is read locally;
images always go through OCR. The result records *how* the text was obtained so
downstream code and audits know whether OCR was involved.

The ordering here is the point. Every page is read before anything is judged,
and every page is judged before anything is uploaded:

    read the whole document locally
        -> classify each page from its own content
            -> resume pages   -> the Veris resume endpoint
               Aadhaar pages  -> the Aadhaar endpoint      (multipass)
               Indian passport pages -> the passport endpoint (multipass)
               everything else -> never uploaded at all

A file therefore cannot be billed as a resume extraction because of what it was
called, and an Aadhaar card on the last page of a forty-page bundle cannot be
missed because the CV was found on page three.

"""
from __future__ import annotations

import re

from app.config import settings
from app.core.exceptions import TextExtractionError, UnsupportedFileTypeError
from app.core.models import ExtractedDocument, PageText
from app.extraction import file_type as ft
from app.extraction import local_ocr
from app.extraction import page_classifier as pc
from app.extraction import pdf_pages
from app.extraction import resume_nationality as rn
from app.extraction.ocr import ocr_via_veris_pages, ocr_via_veris_read
from app.logging_config import get_logger

log = get_logger(__name__)


def extract_text(data: bytes, filename: str = "") -> ExtractedDocument:
    kind = ft.detect(data, filename)
    log.info("Extracting text from %s (category=%s)", filename or "<attachment>", kind.category)

    if kind.category == ft.CATEGORY_PDF:
        return _extract_pdf(data, filename)
    if kind.category == ft.CATEGORY_DOCX:
        return _extract_docx(data)
    if kind.category == ft.CATEGORY_DOC:
        return _extract_doc(data)
    if kind.category == ft.CATEGORY_ODT:
        return _extract_odt(data)
    if kind.category == ft.CATEGORY_IMAGE:
        return _extract_image(data, filename)
    if kind.category == ft.CATEGORY_RTF:
        return _extract_rtf(data)
    if kind.category == ft.CATEGORY_TEXT:
        text = data.decode("utf-8", errors="replace")
        return _classified(_split_pages(text), method="plain")

    raise UnsupportedFileTypeError(f"Cannot extract text from category={kind.category}")


# --------------------------------------------------------------------------- #
def _page_layout_text(page) -> str:
    """One PDF page's text, with a two-column layout un-interleaved."""
    # Group blocks into visual columns (left sidebar vs right main body) to
    # prevent interleaving.
    blocks = page.get_text("blocks")
    mid_x = page.rect.width * 0.4

    valid_blocks = [b for b in blocks if len(b) > 4 and b[4].strip()]
    left_blocks = [b for b in valid_blocks if b[0] < mid_x]
    right_blocks = [b for b in valid_blocks if b[0] >= mid_x]

    left_blocks.sort(key=lambda b: b[1])
    right_blocks.sort(key=lambda b: b[1])

    if left_blocks and right_blocks and len(left_blocks) >= 2 and len(right_blocks) >= 2:
        sorted_blocks = left_blocks + right_blocks
    else:
        sorted_blocks = sorted(valid_blocks, key=lambda b: (b[1], b[0]))

    return "\n\n".join(b[4].strip() for b in sorted_blocks)


_SPACED_CHARS = re.compile(r"\b[A-Za-z]\s[A-Za-z]\s[A-Za-z]\s[A-Za-z]\b")


def _needs_ocr(text: str) -> bool:
    """A page whose text layer is missing, or is per-character garbage."""
    return len(text.strip()) < settings.ocr_min_text_chars or bool(_SPACED_CHARS.search(text))


def _extract_pdf(data: bytes, filename: str = "") -> ExtractedDocument:
    """Read every page of the PDF, then say what is on each of them.

    The text layer is free and is used wherever it is trustworthy; every other
    page is read locally. No page is skipped on the grounds that the résumé has
    already been found — the identity documents in an application bundle are
    usually *behind* the CV, and stopping at the CV is precisely why they were
    never extracted.
    """
    import fitz  # PyMuPDF

    page_texts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        page_count = doc.page_count
        for page in doc:
            page_texts.append(_page_layout_text(page))

    method = "pdf_text"

    # Which pages the embedded text layer cannot answer for. A page with no text
    # at all is a scan; one whose text comes back per-character spaced is a
    # broken font map, and both need to be read as pictures.
    targets = [n for n in range(1, page_count + 1) if _needs_ocr(page_texts[n - 1])]

    # A safety net against a mis-sent thousand-page scan, not a budget: local
    # reading costs CPU, not money. Enforced here rather than only inside the
    # reader so the limit holds however the pages are read, and it is always
    # reported with the page numbers that went unread.
    ceiling = max(1, settings.ocr_max_pages)
    beyond_ceiling: list[int] = []
    if len(targets) > ceiling:
        beyond_ceiling = targets[ceiling:]
        log.warning(
            "'%s' needs %d page(s) read but the safety ceiling is %d; pages %s were "
            "not read locally (raise OCR_MAX_PAGES)",
            filename or "attachment", len(targets), ceiling, beyond_ceiling,
        )
        targets = targets[:ceiling]

    if targets:
        log.info(
            "Reading %d of %d page(s) of '%s' locally (no usable text layer)",
            len(targets), page_count, filename or "attachment",
        )
        confirm = targets
        triage_texts: dict[int, str] = {}

        if settings.ocr_triage_enabled and len(targets) > 1:
            # Stage one: one cheap pass, to find out where the money should go.
            triaged = local_ocr.triage_pdf_pages(
                data, pages=set(targets), filename=filename
            )
            triage_texts = {n: r.text for n, r in triaged.items()}
            confirm = _pages_worth_confirming(targets, triage_texts, filename)

        # Stage two: the real read, on the pages that earned it.
        reads = local_ocr.ocr_pdf_page_reads(data, pages=set(confirm), filename=filename)
        fresh = {number: read.text for number, read in reads.items()}

        # The confirmed text wins wherever it exists; triage text stands for the
        # rest. A page nobody confirmed was one triage read and understood — it
        # is a certificate or an experience letter, and its text is still part
        # of the document, still stored, still searchable. It simply did not
        # need to be read twice.
        for number, text in triage_texts.items():
            fresh.setdefault(number, text)

        page_texts, ocr_pages = _merge_pages(page_texts, fresh)

        # Whatever the local pass could not read, the cloud reader is asked for.
        #
        # This is the difference between "there is nothing on this page" and "we
        # failed to read this page", and until now the pipeline could not tell
        # them apart: both arrived at the classifier as an empty string, scored
        # nothing, and were filed as an ignored certificate. An identity
        # document only reaches the endpoint that can extract it if the local
        # read was already good enough to name it, so every page lost here was
        # lost silently and permanently.
        # Pages the ceiling cut belong here too. The ceiling bounds *local CPU*,
        # which is a reason not to read a page here — not a reason to publish a
        # bundle with a hole in it. Whichever way a page ended up with no text,
        # the question the classifier is about to ask of it is the same one, and
        # it cannot be answered from an empty string. `veris_recover_max_pages`
        # is what bounds the cost of saying so.
        unread = sorted(
            set(n for n in targets if not page_texts[n - 1].strip()) | set(beyond_ceiling)
        )
        # And the pages that did come back with something, but with something
        # the classifier cannot use. `_degraded_reads` explains what that means
        # and why an empty page was never the only way to lose a passport.
        degraded = sorted(_degraded_reads(reads) - set(unread))
        recovered: set[int] = set()
        if unread or degraded:
            page_texts, recovered = _recover_unread_pages(
                data, filename, page_texts, unread, degraded
            )
            ocr_pages |= recovered

        # A hint of an identity document earns a second, closer look.
        #
        # After the recovery above, not before it. The second look is a local
        # re-render at `ocr_deep_read_dpi`, and it is worth spending only on a
        # page the local reader *did* read — one whose caption was too small to
        # resolve, not one it ran out of time on. Asking a host that could not
        # finish a page at 300 DPI to try it again at 450 is 2.25x the pixels
        # for the same empty answer, which is the reasoning `_read_pdf_page`
        # already applies one layer down. So pages the cloud has now answered
        # are excluded: they carry the better read of the two, and re-reading
        # them here could only cost time.
        page_texts, deepened = _deepen_identity_hints(
            data, filename, page_texts, ocr_pages - recovered
        )
        ocr_pages |= deepened

        if ocr_pages:
            method = "pdf_ocr"
    else:
        ocr_pages = set()

    return _classified(
        page_texts,
        method=method,
        page_count=page_count,
        ocr_pages=ocr_pages,
        data=data,
        filename=filename,
    )


def _merge_pages(
    page_texts: list[str], fresh: dict[int, str],
) -> tuple[list[str], set[int]]:
    """Fold freshly-read pages into the text-layer read, keeping page numbers."""
    merged = list(page_texts)
    read: set[int] = set()
    for number, text in fresh.items():
        if not text.strip():
            continue
        while len(merged) < number:
            merged.append("")
        merged[number - 1] = text
        read.add(number)
    return merged, read


#: How far below `ocr_page_quality_floor` a read has to fall before it is
#: treated as a failure rather than as a thin page.
#:
#: The floor itself is the wrong line to draw this at, and drawing it there was
#: the first thing tried. `text_quality` is words x readable-character density,
#: and identity documents are *sparse by design*: an Aadhaar front carries maybe
#: fifteen words and scores about 11 against a floor of 12. Recovering every
#: page under the floor therefore means uploading every ID card in every bundle
#: — the exact documents the local reader handled correctly.
#:
#: What separates a sparse page from a failed one is not how few words there are
#: but how much of the page is not words at all. That same Aadhaar front scores
#: ~11; a half-size rescue of a passport that produced eighty characters of
#: speckle scores under 2. A quarter of the floor (3.0 by default) sits in the
#: gap with room on both sides, and it moves with the floor rather than being a
#: second number to keep in step with it.
_DEGRADED_QUALITY_FRACTION = 0.25


def _degraded_reads(reads: "dict[int, local_ocr.PageRead]") -> set[int]:
    """Pages that came back with text the classifier cannot actually use.

    An empty page is the *visible* way local OCR loses a document, and it is the
    one the recovery below was built for. It is not the common one.

    What actually happens on a starved host is subtler and worse. A full-size
    pass times out, `local_ocr` retries the page at half width, and half a scan
    of a passport answers with a few dozen characters of speckle — which is how
    a real bundle produced ``Page 3 read as empty at 2550px but answered at
    1600px: 81 chars recovered``, and then reported no passport. Those 81
    characters are not a read. They are a failure wearing the shape of one, and
    every gate downstream is sized for text rather than for noise:

    * under `pc.too_short_to_classify` the classifier scores the page at zero by
      construction — it is not judged badly, it is not judged;
    * `has_identity_hint` needs to see a word like "passport" or "nationality"
      survive the read, which is precisely what a garbled page does not offer,
      so the second look never fires either.

    So a degraded read is worse than an empty one: it fails the same way *and*
    it suppresses the recovery an empty page would have received. Three
    signatures name it, and each is something the reader already knew and threw
    away when it handed back nothing but a string:

    * `tesseract:rescue` — the page only answered after being shrunk. The reader
      itself calls this the last rung, taken when a read did not finish;
    * too short for `page_classifier` to form any verdict about;
    * quality far enough under the floor to be noise rather than a sparse page
      (see `_DEGRADED_QUALITY_FRACTION`).

    Cheap to act on, too: these pages join the *same* subset PDF and the same
    single request as the unread ones, so a bundle with two empty pages and two
    degraded ones costs one call either way — a larger upload, not a second
    extraction.
    """
    if not settings.veris_recover_degraded_pages:
        return set()

    noise_ceiling = float(settings.ocr_page_quality_floor) * _DEGRADED_QUALITY_FRACTION
    degraded: set[int] = set()
    for number, read in reads.items():
        text = read.text or ""
        if not text.strip():
            continue  # empty; already counted as unread by the caller
        if (
            read.engine == "tesseract:rescue"
            or pc.too_short_to_classify(text)
            or read.quality < noise_ceiling
        ):
            degraded.add(number)
    return degraded


def _deepen_identity_hints(
    data: bytes, filename: str, page_texts: list[str], ocr_pages: set[int],
) -> "tuple[list[str], set[int]]":
    """Read again, harder, any page that hinted at an identity document.

    The scores this keys off are deliberately coarse — a strong marker is worth
    2.0, a weak one 0.5, and 3.0 routes the page. The consequence is that a page
    carrying a single half-read trace of a passport scores 0.5 and is dropped
    without anything ever having looked at it properly. That is backwards. A
    hint is the signal to spend *more* effort on a page, not the signal to stop:
    the page either firms up into a document, or it settles into nothing, and
    either way the bundle stops guessing.

    Concretely, this is the difference between a passport's back page reading
    "Name of Father / Legal Guardian" — one marker, 2.0, under the seed — and
    the same page at 450 DPI picking up the second caption that carries it over.

    Only pages below the seed are re-read: one that already routes has nothing
    to gain. Only the better read is kept, judged on the identity evidence it
    produced rather than on length, because a longer read full of scanner
    speckle is not a better one.
    """
    if not (settings.ocr_deep_read_enabled and ocr_pages):
        return page_texts, set()

    def routable_evidence(text: str) -> float:
        """The best score among the kinds that can actually be sent somewhere.

        Not `max(...)` over every kind: the generic `document` score is scored
        but never routed (`is_document` is switched off in the classifier), so
        including it sends pages for an expensive second read on the strength of
        evidence that cannot change any outcome. Four such pages in one bundle
        were re-read at 450 DPI for nothing.
        """
        scores = pc.id_document_scores(text)
        return max(scores.get(pc.AADHAAR, 0.0), scores.get(pc.PASSPORT, 0.0))

    # Any trace at all, scored or not.
    #
    # A score of 0.0 does not mean "nothing here" — `id_document_scores` only
    # counts markers precise enough to be evidence, so a page where OCR half-read
    # the caption keeps the bare word "passport" and still scores zero. Those are
    # exactly the pages worth looking at again, so `has_identity_hint` brings
    # them in on a much looser test than the one used to route anything.
    #
    # A page already over the seed is left alone: it routes, and re-reading it
    # could only change an answer that is already correct.
    floor = float(settings.ocr_deep_read_score)
    candidates: list[tuple[float, int]] = []
    for number in sorted(ocr_pages):
        text = page_texts[number - 1]
        best = routable_evidence(text)
        if best >= pc.ID_SEED_SCORE:
            continue
        if best >= floor and best > 0:
            candidates.append((best, number))
        elif pc.has_identity_hint(text):
            candidates.append((best, number))

    if not candidates:
        return page_texts, set()

    # Strongest hints first: if the ceiling bites, it should bite on the pages
    # least likely to have been a document in the first place.
    candidates.sort(reverse=True)
    ceiling = max(1, int(settings.ocr_deep_read_max_pages))
    wanted = [n for _score, n in candidates[:ceiling]]
    if len(candidates) > ceiling:
        log.warning(
            "'%s' has %d page(s) hinting at an identity document but only %d may "
            "be re-read; page(s) %s keep their first read (raise "
            "OCR_DEEP_READ_MAX_PAGES)",
            filename or "attachment", len(candidates), ceiling,
            [n for _s, n in candidates[ceiling:]],
        )

    log.info(
        "Second look at page(s) %s of '%s' at %ddpi: each shows a trace of an "
        "identity document but scores under the %.1f needed to route it",
        wanted, filename or "attachment", settings.ocr_deep_read_dpi,
        pc.ID_SEED_SCORE,
    )
    try:
        # One pass at high resolution, not the whole ladder. The ladder exists
        # to try different *segmentations* of a page that read poorly; this page
        # read fine, it simply rendered a caption too small to resolve, and the
        # only thing that fixes that is pixels. Running psm 4 and psm 3 as well
        # would triple the cost of a step that has to be cheap enough to run on
        # every hint rather than only on strong ones.
        reads = local_ocr.single_pass_pages(
            data, set(wanted), filename,
            dpi=int(settings.ocr_deep_read_dpi), label="second look",
        )
    except Exception as exc:  # noqa: BLE001 — the first read still stands
        log.warning("Second look at '%s' failed (%s); keeping the first read", filename, exc)
        return page_texts, set()

    merged = list(page_texts)
    improved: set[int] = set()
    for number, read in reads.items():
        text = read.text or ""
        if not text.strip():
            continue
        before = routable_evidence(merged[number - 1])
        after = routable_evidence(text)
        if after <= before:
            continue
        merged[number - 1] = text
        improved.add(number)
        log.info(
            "Page %d of '%s' read again at %ddpi: identity evidence %.1f -> %.1f%s",
            number, filename or "attachment", settings.ocr_deep_read_dpi,
            before, after,
            " (now routable)" if after >= pc.ID_SEED_SCORE else "",
        )

    if not improved:
        log.info(
            "The second look at page(s) %s of '%s' found nothing more; they are "
            "not identity documents",
            wanted, filename or "attachment",
        )
    return merged, improved


def _pages_worth_confirming(
    targets: list[int], triage: dict[int, str], filename: str = "",
) -> list[int]:
    """Which pages the cheap pass says are worth a full-resolution read.

    Recall, not precision. A page wrongly nominated costs one read; a page
    wrongly dropped costs a passport, and that is not a trade to make finely.
    So the bar is deliberately on the floor: *any* identity signal at all, any
    hint of a résumé, and anything the triage could not read confidently enough
    to have an opinion about.

    That last clause is doing most of the work, and it is why this is safe
    despite the classifier not being stable across resolution. Measured on the
    Saravanan bundle, a 200-DPI read loses the passport on page 27 outright —
    but it does not read page 27 as *nothing*, it reads it as a page with too
    little on it to judge, which lands here and is confirmed at 300 DPI where
    the passport is found. Triage is never allowed to be the last word on a
    page; it only decides what gets looked at properly.
    """
    floor = max(0, int(settings.ocr_triage_min_chars))
    confirm: list[int] = []
    for number in targets:
        text = triage.get(number, "")
        stripped = text.strip()
        if not stripped or len(stripped) < floor:
            confirm.append(number)          # too little read to have an opinion
            continue
        if max(pc.id_document_scores(text).values() or [0.0]) > 0:
            confirm.append(number)          # any identity signal whatsoever
            continue
        page = pc.classify_page(text, number)
        if page.kind in (pc.RESUME, pc.UNKNOWN):
            confirm.append(number)          # a résumé, or nothing conclusive

    log.info(
        "Triage of '%s': %d of %d page(s) go on to a full read (%s)",
        filename or "attachment", len(confirm), len(targets),
        confirm if len(confirm) <= 20 else f"{confirm[:20]}...",
    )
    return confirm


def _recover_unread_pages(
    data: bytes, filename: str, page_texts: list[str],
    unread: list[int], degraded: "list[int] | None" = None,
) -> "tuple[list[str], set[int]]":
    """Re-read the pages local OCR lost, through Veris.

    Two ways a page is lost, one call to fetch them both:

    * **unread** — nothing came back at all. Under CPU pressure a Tesseract pass
      hits `ocr_page_timeout_seconds` and returns "", and by the time that
      reaches the classifier it is indistinguishable from a blank sheet.
    * **degraded** — something came back, but not something that can be
      classified: a half-size rescue read, a page too short to score, or noise.
      See `_degraded_reads`. These are the worse half, because a page with a
      little text on it does not look lost to anything downstream.

    Only those pages, and only when there are some: a bundle that read cleanly
    never gets here, so this costs nothing on the common case.

    The pages are sent as a subset PDF for the same reason the resume and
    identity payloads are — a trimmed page keeps the scanner's full resolution,
    and the upload is most of the round trip — and the result is merged back at
    the page numbers it came from, so the classifier that runs next sees one
    complete document rather than a local read with holes in it.

    An unread page takes whatever the cloud reader returns; there is nothing to
    weigh it against. A degraded page keeps its local text unless the cloud read
    genuinely beats it on `text_quality` — the same judgement `local_ocr` uses
    when two of its own passes disagree, and the reason recovery cannot make a
    page worse than it already was.

    A page the cloud reader cannot read either stays as it was. That is a real
    finding rather than a race, and it is logged as one.
    """
    degraded = list(degraded or [])
    if not (settings.veris_recover_unread_pages and settings.veris_ocr_api_key):
        log.warning(
            "Local OCR could not usably read page(s) %s of '%s' (unread: %s; "
            "degraded: %s) and cloud recovery is off; anything on those pages — "
            "an identity document included — cannot be classified (set "
            "VERIS_OCR_API_KEY and VERIS_RECOVER_UNREAD_PAGES=true)",
            sorted(set(unread) | set(degraded)), filename or "attachment",
            unread, degraded,
        )
        return page_texts, set()

    degraded_pages = set(degraded)
    unread_pages = set(unread)
    lost = sorted(unread_pages | degraded_pages)
    ceiling = max(1, settings.veris_recover_max_pages)
    wanted = lost
    if len(lost) > ceiling:
        # Unread before degraded when the ceiling bites. A page with no text at
        # all has strictly less to lose than one holding a poor read, so if only
        # some can be bought back, buy the ones that are entirely gone.
        ordered = sorted(unread_pages) + sorted(degraded_pages - unread_pages)
        wanted = sorted(ordered[:ceiling])
        log.warning(
            "'%s' has %d page(s) local OCR could not usably read but only %d may "
            "be recovered; page(s) %s keep what they have (raise "
            "VERIS_RECOVER_MAX_PAGES)",
            filename or "attachment", len(lost), ceiling, sorted(ordered[ceiling:]),
        )

    subset = pdf_pages.subset_pdf(data, wanted)
    if subset is None and len(wanted) >= len(page_texts):
        # Every page failed. `subset_pdf` declines to build a "subset" that is
        # the whole document, and the right answer to that is the whole
        # document — not to abandon a bundle precisely because none of it read.
        subset = data
    if not subset:
        log.warning(
            "Could not isolate page(s) %s of '%s' for recovery",
            wanted, filename or "attachment",
        )
        return page_texts, set()

    log.info(
        "Local OCR lost page(s) %s of '%s' (%d unread, %d read too poorly to "
        "classify); asking Veris for them",
        wanted, filename or "attachment",
        sum(1 for n in wanted if n not in degraded_pages),
        sum(1 for n in wanted if n in degraded_pages),
    )
    payload = pdf_pages.compact_pdf(
        subset, settings.ocr_payload_max_bytes, settings.ocr_payload_dpi
    )
    try:
        read = ocr_via_veris_read(payload, filename or "attachment.pdf")
    except Exception as exc:  # noqa: BLE001 — the local read still stands
        log.warning(
            "Veris recovery of page(s) %s of '%s' failed (%s); those pages keep "
            "what the local read gave them",
            wanted, filename or "attachment", exc,
        )
        return page_texts, set()

    merged = list(page_texts)
    recovered: set[int] = set()
    kept_local: list[int] = []
    for index, number in enumerate(wanted):
        text = read.pages[index] if index < len(read.pages) else ""
        if not (text or "").strip():
            continue
        if number in degraded_pages and local_ocr.text_quality(text) <= local_ocr.text_quality(
            merged[number - 1]
        ):
            # The page already had a poor read and this one is no better. Taking
            # it anyway would be trading one unusable page for another while
            # discarding the only evidence in hand.
            kept_local.append(number)
            continue
        merged[number - 1] = text
        recovered.add(number)

    still_lost = [n for n in wanted if n not in recovered and n not in kept_local]
    if recovered:
        log.info(
            "Veris recovered page(s) %s of '%s' (%d chars) that local OCR could not read",
            sorted(recovered), filename or "attachment",
            sum(len(merged[n - 1]) for n in recovered),
        )
    if kept_local:
        log.info(
            "Page(s) %s of '%s' read no better at Veris than locally; keeping the "
            "local text",
            kept_local, filename or "attachment",
        )
    if still_lost:
        log.warning(
            "Page(s) %s of '%s' read as empty both locally and at Veris; treating "
            "them as genuinely blank",
            still_lost, filename or "attachment",
        )
    return merged, recovered


def _refine_resume_pages(
    data: bytes, filename: str, page_texts: list[str], resume_pages: list[int],
) -> "tuple[list[str], dict | None]":
    """Re-read the résumé pages — and only those — through the Veris endpoint.

    This is the one place anything is uploaded, and it runs *after* the local
    read has established that these pages are a résumé. That ordering is the
    fix for the original complaint: a bank statement or a job-board digest is
    now identified locally and never reaches a paid extraction, while a real CV
    still gets the better read on the two pages that hold it.
    """
    if not (settings.veris_refine_resume_pages and settings.veris_ocr_api_key):
        return page_texts, None
    if not resume_pages:
        return page_texts, None

    # Page trim, then size trim — the same treatment the Aadhaar and passport
    # payloads get, for the same reason: a trimmed page still carries the
    # scanner's full resolution, and that upload is most of the round trip.
    payload = pdf_pages.compact_pdf(
        pdf_pages.subset_pdf(data, resume_pages) or data,
        settings.ocr_payload_max_bytes,
        settings.ocr_payload_dpi,
    )
    try:
        read = ocr_via_veris_read(payload, filename or "resume.pdf")
        texts = [t or "" for t in read.pages]
        extraction = read.result
    except Exception as exc:  # noqa: BLE001 — the local read already stands
        log.warning("Veris refinement of pages %s failed (%s); keeping the local read",
                    resume_pages, exc)
        return page_texts, None

    if not any(t.strip() for t in texts):
        return page_texts, extraction

    refined = list(page_texts)
    if len(texts) == len(resume_pages):
        for number, text in zip(resume_pages, texts):
            if len(text.strip()) > len(refined[number - 1].strip()):
                refined[number - 1] = text
    else:
        # Veris merged or dropped a page, so the per-page mapping cannot be
        # trusted. The combined text goes on the first résumé page: the boundary
        # is lost, the content is not.
        combined = "\n\n".join(t for t in texts if t.strip())
        if len(combined.strip()) > len(refined[resume_pages[0] - 1].strip()):
            refined[resume_pages[0] - 1] = combined
    log.info("Refined résumé page(s) %s of '%s' through Veris", resume_pages, filename)
    return refined, extraction



def _classified(
    page_texts: list[str],
    method: str,
    page_count: int | None = None,
    ocr_pages: set[int] | None = None,
    data: bytes | None = None,
    filename: str = "",
) -> ExtractedDocument:
    """Wrap extracted page text with the classifier's verdict.

    ``text`` stays the full extraction — every page, nothing dropped — and the
    classification records which slice of it is the candidate's profile.
    """
    ocr_pages = ocr_pages or set()

    # "I read it and it is not a resume" and "I could not read it" are opposite
    # findings, and collapsing them is dangerous: the first is a permanent skip,
    # the second is a retryable failure. A scanned 9-page CV with Tesseract
    # missing produced *zero* characters, which the classifier — correctly, on
    # the evidence it had — scored as a non-resume. That would have rejected
    # every scanned application in the mailbox as a misnamed file, and marked
    # each one handled. An empty extraction is an extraction error.
    if not any((t or "").strip() for t in page_texts):
        raise TextExtractionError(
            f"No text could be extracted from this {method} document "
            f"({len(page_texts)} page(s)). It is most likely a scan with no text "
            "layer while OCR is unavailable: install Tesseract (or set "
            "TESSERACT_CMD) or configure VERIS_OCR_API_KEY."
        )

    result = pc.classify_document(page_texts)

    # Whose CV is this? Asked here, of the local read, and before the upload
    # below — a candidate this desk cannot place must not cost a résumé
    # extraction. The classifier has already run the passport reader over every
    # booklet page in the bundle, so an MRZ naming the issuing state is evidence
    # already in hand rather than evidence to go and fetch.
    #
    # The answer is carried on the result, not recomputed by the pipeline: it
    # has to hold back the upload *and* keep the record out of the database, and
    # one policy evaluated twice is two chances to disagree with itself.
    whole_text = "\n\n".join(t for t in page_texts if (t or "").strip())
    nationality = rn.detect_resume_nationality(
        whole_text,
        passport_verdicts=rn.passport_verdicts_from(page_texts, result.page_kinds),
    )
    accepted, nationality_reason = rn.should_ingest(nationality)
    if not accepted:
        log.info("Nationality filter: %s", nationality_reason)

    # Only now — with the whole document read, the résumé located and the
    # candidate's nationality established — is anything sent out for a better
    # read, and only the pages that hold it.
    veris_resume_result: dict | None = None
    if (
        accepted
        and data is not None
        and result.is_resume
        and result.resume_pages
        and ocr_pages
    ):
        refined, veris_resume_result = _refine_resume_pages(
            data, filename, page_texts, result.resume_pages
        )
        if refined is not page_texts:
            page_texts = refined
            result = pc.classify_document(page_texts)

    pages = [
        PageText(
            page_number=p.page_number,
            text=page_texts[p.page_number - 1],
            kind=p.kind,
            score=p.score,
            ocr_used=p.page_number in ocr_pages,
        )
        for p in result.pages
    ]
    text = "\n\n".join(t.strip() for t in page_texts if t.strip())
    log.info(
        "\n================================================================================\n"
        "[ENDPOINT: RESUME INGESTED]\n"
        "  * File: %s (Total Pages: %d | OCR Pages: %s)\n"
        "  * Classification: is_resume=%s (Confidence: %.2f)\n"
        "  * Resume Pages: %s\n"
        "  * Candidate Nationality: %s (Accepted: %s)\n"
        "  * Reason: %s\n"
        "================================================================================",
        filename or "attachment", page_count if page_count is not None else len(page_texts),
        sorted(ocr_pages) if ocr_pages else "None", result.is_resume, result.confidence,
        result.resume_pages or "None", nationality.country or "IND", accepted, result.reason,
    )
    return ExtractedDocument(
        text=text,
        method=method,
        page_count=page_count if page_count is not None else len(page_texts),
        ocr_used=bool(ocr_pages),
        char_count=len(text),
        pages=pages,
        resume_pages=result.resume_pages,
        is_resume=result.is_resume,
        classification_confidence=result.confidence,
        classification_reason=result.reason,
        veris_resume_result=veris_resume_result,
        nationality=nationality.as_dict(),
        nationality_accepted=accepted,
        nationality_reason=nationality_reason,
    )


def _extract_docx(data: bytes) -> ExtractedDocument:
    import io

    # Primary: python-docx (captures paragraphs + tables).
    try:
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(parts).strip()
        if text:
            return _classified(_split_pages(text), method="docx")
    except Exception as exc:  # noqa: BLE001 — fall through to docx2txt
        log.warning("python-docx failed (%s); trying docx2txt", exc)

    # Fallback: docx2txt.
    try:
        import docx2txt

        with io.BytesIO(data) as buf:
            text = (docx2txt.process(buf) or "").strip()
        if text:
            return _classified(_split_pages(text), method="docx")
    except Exception as exc:  # noqa: BLE001
        raise TextExtractionError(f"Could not read .docx: {exc}") from exc

    raise TextExtractionError("DOCX contained no extractable text.")


def _extract_doc(data: bytes) -> ExtractedDocument:
    """Legacy binary .doc — needs an external converter (LibreOffice/antiword).

    We attempt a headless LibreOffice conversion to .docx if available, then reuse
    the docx path. If no converter is present we raise a clear, actionable error.
    """
    import shutil
    import subprocess
    import tempfile
    from pathlib import Path

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise UnsupportedFileTypeError(
            "Legacy .doc files require LibreOffice for conversion. Install LibreOffice "
            "(soffice on PATH), or ask senders to submit .docx/.pdf."
        )
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "resume.doc"
        src.write_bytes(data)
        subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, str(src)],
            check=True, capture_output=True, timeout=120,
        )
        converted = Path(tmp) / "resume.docx"
        if not converted.exists():
            raise TextExtractionError("LibreOffice did not produce a .docx from the .doc.")
        result = _extract_docx(converted.read_bytes())
        result.method = "doc"
        return result


def _extract_odt(data: bytes) -> ExtractedDocument:
    """OpenDocument text — a zip whose `content.xml` holds the paragraphs.

    Worth the twenty lines: LibreOffice exports .odt by default, so it is what
    arrives from candidates who do not own Word, and the alternative is telling
    them their application was unreadable.
    """
    import io
    import zipfile
    from xml.etree import ElementTree

    _TEXT_NS = "urn:oasis:names:tc:opendocument:xmlns:text:1.0"

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            xml = archive.read("content.xml")
    except (zipfile.BadZipFile, KeyError) as exc:
        raise TextExtractionError(f"Could not read .odt archive: {exc}") from exc

    try:
        root = ElementTree.fromstring(xml)
    except ElementTree.ParseError as exc:
        raise TextExtractionError(f"Could not parse .odt content.xml: {exc}") from exc

    lines: list[str] = []
    # Paragraphs and headings, in document order. `itertext` on each keeps the
    # spans inside a paragraph (bold runs, hyperlinks) joined to their line.
    for node in root.iter():
        if node.tag in (f"{{{_TEXT_NS}}}p", f"{{{_TEXT_NS}}}h"):
            line = "".join(node.itertext()).strip()
            if line:
                lines.append(line)

    text = "\n".join(lines).strip()
    if not text:
        raise TextExtractionError("ODT contained no extractable text.")
    return _classified(_split_pages(text), method="odt")


def _extract_image(data: bytes, filename: str = "") -> ExtractedDocument:
    """A photographed or scanned page, read locally before it is judged.

    An image attachment carries no filename evidence worth anything and no text
    layer to inspect, so it used to be uploaded to find out what it was — which
    meant every signature graphic and every marketing banner that cleared the
    size floor was billed as a résumé extraction. It is read here first; only if
    the content reads as a résumé does `_classified` send it anywhere.
    """
    text = local_ocr.ocr_image_bytes(data).strip()
    if len(text) < settings.ocr_min_text_chars:
        log.warning("Image OCR produced only %d chars for '%s'", len(text), filename)
    if not text:
        raise TextExtractionError("Image OCR produced no text.")
    return _classified([text], method="image_ocr", page_count=1, ocr_pages={1})


def _extract_rtf(data: bytes) -> ExtractedDocument:
    # Minimal RTF → text: strip control words. Good enough to feed the AI.
    raw = data.decode("latin-1", errors="replace")
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", raw)
    text = re.sub(r"[{}]", " ", raw)
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        raise TextExtractionError("RTF produced no text.")
    return _classified([text], method="rtf", page_count=1)


def _split_pages(text: str) -> list[str]:
    """Page-break the formats that have no page model of their own.

    Word and plain-text documents carry no page count, but a converted or
    concatenated bundle usually still has form feeds where the pages were.
    """
    pages = [part for part in (text or "").split("\f")]
    return pages if len(pages) > 1 else [text or ""]
